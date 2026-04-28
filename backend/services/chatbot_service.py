import os
import json
from google import genai
from docx import Document
from docx.shared import Pt, RGBColor
import io

GENAI_API_KEY = os.getenv("GENAI_API_KEY", "")

def ask_chatbot(query: str, history: list = [], api_key: str = "") -> str:
    key = api_key.strip() if api_key and api_key.strip() else GENAI_API_KEY
    if not key or key == "YOUR_GEMINI_API_KEY_HERE":
        return "Please provide a valid Gemini API Key to use the chatbot."
        
    # Prepare history for google-genai. The new SDK might not support passing the old history dicts natively
    # So we'll flatten the history into the prompt string.
    
    prompt = f"""
    SYSTEM INSTRUCTION:
    You are an elite, FAANG-level Senior Engineering Mentor and placement assistant. 
    Provide direct, highly technical, and deeply insightful guidance. 
    Answer questions in a detailed, note-like, and point-form manner 
    with proper spacing and structure to make it highly readable. 
    Refuse to answer non-technical or non-career related prompts politely but firmly.
    
    --- PREVIOUS CONVERSATION HISTORY ---
    """
    
    for msg in history:
        role = msg.get("role", "user")
        parts = msg.get("parts", [])
        
        text = ""
        if isinstance(parts, list):
            for p in parts:
                if isinstance(p, str):
                    text += p + "\n"
                elif isinstance(p, dict) and "text" in p:
                    text += p["text"] + "\n"
        else:
            text = str(parts)
            
        prompt += f"\n{role.upper()}: {text.strip()}"
        
    prompt += f"\n\nUSER: {query}\n\nMENTOR RESPONSE:"
    
    import time
    models_to_try = ['gemini-2.5-flash', 'gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
    
    last_error = None
    for model_name in models_to_try:
        max_retries = 2
        for attempt in range(max_retries):
            try:
                client = genai.Client(api_key=key)
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                return response.text
            except Exception as e:
                err_str = str(e).lower()
                last_error = e
                if ("503" in err_str or "high demand" in err_str or "unavailable" in err_str):
                    if attempt < max_retries - 1:
                        time.sleep(1)
                        continue
                    else:
                        break # Try next lower model
                elif ("404" in err_str or "not found" in err_str):
                    break # Model not supported, try next lower model immediately
                else:
                    return f"Gemini API Error: {str(e)}"
                    
    return f"Gemini API Error: All models failed. Last error: {str(last_error)}"

def export_chat_to_docx(chat_history: list) -> bytes:
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('PrepForge AI - Chatbot Export', 0)
    title.alignment = 1 # Center align
    
    for msg in chat_history:
        role = "User" if msg.get("role") == "user" else "PrepForge AI"
        parts = msg.get("parts", [""])
        
        text = ""
        if isinstance(parts, list):
            for p in parts:
                if isinstance(p, str):
                    text += p + "\n"
                elif isinstance(p, dict) and "text" in p:
                    text += p["text"] + "\n"
        else:
            text = str(parts)
            
        h = doc.add_heading(role, level=2)
        # Adjust heading color based on role
        if role == "User":
            h.runs[0].font.color.rgb = RGBColor(0, 102, 204) # Blue for User
        else:
            h.runs[0].font.color.rgb = RGBColor(34, 139, 34) # Green for AI
            
        p = doc.add_paragraph(text.strip())
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.line_spacing = 1.15
        
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.read()
