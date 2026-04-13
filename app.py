import os
import json
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
import PyPDF2
import traceback
from io import BytesIO
from docx import Document
from flask import send_file

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5 MB max

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def load_data(filename):
    filepath = os.path.join(app.root_path, 'data', filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/cs-subjects')
def cs_subjects():
    topics = load_data('cs_topics.json')
    return render_template('cs_subjects.html', topics=topics)

@app.route('/cs-subjects/<subject_id>')
def subject_detail(subject_id):
    topics = load_data('cs_topics.json')
    if subject_id not in topics:
        return "Subject not found", 404
        
    subject_data = topics[subject_id]
    return render_template('subject_detail.html', subject=subject_data, subject_id=subject_id)

@app.route('/cs-subjects/<subject_id>/examples')
def subject_examples(subject_id):
    try:
        examples = load_data('examples.json')
    except:
        examples = {}
    if subject_id not in examples:
        return "Examples not found", 404
        
    example_data = examples[subject_id]
    return render_template('examples.html', subject=example_data, subject_id=subject_id)

@app.route('/gd-prep')
def gd_prep():
    gd_data = load_data('gd_topics.json')
    return render_template('gd_prep.html', gd_data=gd_data)

@app.route('/hr-interview', methods=['GET', 'POST'])
def hr_interview():
    hr_data = load_data('hr_questions.json')
    generated_questions = []
    error = None

    if request.method == 'POST':
        github_link = request.form.get('github_link')
        file = request.files.get('resume')
        
        extracted_text = ""
        
        try:
            # Handle github link
            if github_link:
                # Basic mock logic for Github parsing
                extracted_text += f" GitHub Profile: {github_link} "
            
            # Handle PDF resume
            if file and file.filename != '':
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                # Extract text using PyPDF2
                with open(filepath, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text
                            
                # Cleanup
                if os.path.exists(filepath):
                    os.remove(filepath)
                    
            if extracted_text:
                # Mock keyword based generation
                generated_questions = generate_mock_questions(extracted_text)
            else:
                error = "Please provide a GitHub link or upload a PDF resume."
                
        except Exception as e:
            traceback.print_exc()
            error = "An error occurred while processing your input."

    return render_template('hr_interview.html', hr_data=hr_data, generated_questions=generated_questions, error=error)

def generate_mock_questions(text):
    text_lower = text.lower()
    questions = []
    
    if 'github.com' in text_lower:
        questions.extend([
            ("Walk me through the architecture of your most complex GitHub repository.", "Discuss layers, controllers, services, and how they interact. Focus on structural design."),
            ("How did you manage version control and branch merging during this project?", "Mention PRs, git rebase vs merge, dealing with merge conflicts."),
            ("I see this project is open-source. Have you dealt with public issues or community contributions?", "Focus on code review etiquette and setting up Contribution guidelines."),
            ("What security vulnerabilities did you protect against in your codebase?", "Mention sanitization, SQL Injection prevention,, or CORS handling.")
        ])

    if 'python' in text_lower or 'flask' in text_lower or 'django' in text_lower:
        questions.extend([
            ("Why did you choose Python for this backend component over Go or Node?", "Mention rapid prototyping, ecosystem, GIL limitations, and async features."),
            ("How does Python handle garbage collection?", "Explain reference counting and the cyclic garbage collector."),
            ("Explain the concept of Web Server Gateway Interface (WSGI) in your Python apps.", "Discuss how Flask communicates with Gunicorn or the web server.")
        ])
        
    if 'react' in text_lower or 'frontend' in text_lower:
        questions.extend([
            ("Why React? How does the Virtual DOM actually improve performance here?", "Explain batching updates, diffing algorithm (Reconciliation)."),
            ("How did you manage global state in this application?", "Talk about Redux, Context API, or Zustand, and why you avoided prop-drilling."),
            ("Explain a challenging UI bug you resolved involving component lifecycles.", "Mention useEffect dependencies, stale closures, or unmounting memory leaks.")
        ])
        
    if 'sql' in text_lower or 'database' in text_lower:
        questions.extend([
            ("Walk me through your database schema. Did you normalize to 3NF?", "Discuss entities, relations, and trade-offs of redundancy."),
            ("How would you scale this database if traffic spiked by 100x tomorrow?", "Discuss Read Replicas, Caching (Redis), Connection Pooling, or Sharding."),
            ("Explain an advanced SQL query you wrote. Did you use Window Functions or CTEs?", "Provide a concrete narrative of grouping/aggregating complex data.")
        ])

    # Ensure exactly 8-10 high quality questions if keyword matched
    if not questions:
        questions = [
            ("Walk me through your resume, specifically focusing on the timeline of your skills acquisition.", "Connect the dots between your internships and side-projects."),
            ("Describe the most difficult technical hurdle listed on your resume.", "Use STAR method. Emphasize the debugging process."),
            ("What architectural patterns do you usually fall back on when starting a project?", "Discuss Monolith vs Microservices, MVC, etc."),
            ("How do you ensure your code is maintainable for the next developer?", "Mention clean code principles, SOLID, and documentation."),
            ("Tell me about a time you had to pivot a project idea mid-way.", "Highlight adaptability and agile principles."),
            ("Where is the 'weakest link' in the tech stack you listed?", "Show self-awareness of technology limits."),
            ("If you lead a team of 3 juniors on your current project, how would you distribute tasks?", "Focus on pair programming, code reviews, and scoping."),
            ("What is the next language/framework you intend to learn, and why?", "Demonstrate continuous learning.")
        ]
    
    # Cap at 10 and format
    final_questions = questions[:10]
    return [{'question': q, 'hint': h} for q, h in final_questions]
@app.route('/export-hr', methods=['POST'])
def export_hr():
    import json
    questions_str = request.form.get('questions_data')
    if not questions_str:
        return "No data provided", 400
        
    try:
        questions = json.loads(questions_str)
    except:
        return "Invalid data", 400

    doc = Document()
    doc.add_heading('Your Personalized Interview Questions', 0)
    
    for i, item in enumerate(questions, 1):
        q = item.get('question', '')
        h = item.get('hint', '')
        doc.add_heading(f"Q{i}: {q}", level=1)
        doc.add_paragraph(h, style='Intense Quote')
        doc.add_paragraph("") # Space
        
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    return send_file(f, as_attachment=True, download_name='Interview_Prep.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/chatbot')
def chatbot():
    return render_template('chatbot.html')

@app.route('/api/chat', methods=['POST'])
def api_chat():
    data = request.get_json()
    message = data.get('message', '').lower()
    
    # Simple placeholder logic
    response = "I am a simple placeholder AI assistant. I received: " + message
    
    if 'os' in message or 'operating system' in message:
        response = "For OS preparation, focus on Deadlocks, Process Scheduling, and Virtual Memory. Check out the CS Subjects page for more!"
    elif 'dbms' in message or 'database' in message:
        response = "Database normalization, ACID properties, and SQL queries are highly asked in DBMS interviews."
    elif 'gd' in message or 'group discussion' in message:
        response = "In a GD, clarity of thought and confidence are key. Don't just speak frequently, speak thoughtfully. Make sure to let others speak."
    
    return jsonify({"response": response})

@app.route('/export-chat', methods=['POST'])
def export_chat():
    import json
    chat_str = request.form.get('chat_data')
    if not chat_str:
        return "No data provided", 400
        
    try:
        messages = json.loads(chat_str)
    except:
        return "Invalid data", 400

    doc = Document()
    doc.add_heading('PrepForge AI - Chat Transcript', 0)
    
    for item in messages:
        sender = item.get('sender', '')
        text = item.get('text', '')
        doc.add_heading(f"{sender}:", level=2)
        doc.add_paragraph(text)
        doc.add_paragraph("") # Space
        
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    return send_file(f, as_attachment=True, download_name='Chat_Transcript.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True, port=5000)
